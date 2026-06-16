from pathlib import Path
import os
from datetime import datetime
from textwrap import wrap

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\32545\Desktop\软件设计\ruanjiansheji")
DESKTOP = Path(os.environ["USERPROFILE"]) / "Desktop"
OUT = DESKTOP / "智能音箱软件项目六部分报告-按改进方案重生成.docx"
ASSET_DIR = ROOT / "runtime" / "report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
ARCH_IMG = ASSET_DIR / "system_architecture.png"
ER_IMG = ASSET_DIR / "core_er.png"

FONT_PATHS = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\simsun.ttc",
]
FONT_PATH = next((p for p in FONT_PATHS if Path(p).exists()), None)


def font(size, bold=False):
    if bold and Path(r"C:\Windows\Fonts\msyhbd.ttc").exists():
        return ImageFont.truetype(r"C:\Windows\Fonts\msyhbd.ttc", size)
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size)
    return ImageFont.load_default()


def multiline(draw, text, x, y, max_chars, fnt, fill=(32, 32, 32), line_gap=6):
    lines = []
    for part in str(text).split("\n"):
        if len(part) <= max_chars:
            lines.append(part)
        else:
            lines.extend(wrap(part, width=max_chars))
    yy = y
    for line in lines:
        draw.text((x, yy), line, font=fnt, fill=fill)
        yy += fnt.size + line_gap
    return yy


def draw_box(draw, xy, title, lines=(), fill="#F7FAF8", outline="#4F6F64", title_fill="#4F6F64"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw.rounded_rectangle((x1, y1, x2, y1 + 46), radius=18, fill=title_fill, outline=title_fill)
    draw.rectangle((x1, y1 + 28, x2, y1 + 46), fill=title_fill)
    draw.text(((x1 + x2) // 2, y1 + 23), title, font=font(22, True), fill="white", anchor="mm")
    yy = y1 + 62
    for line in lines:
        yy = multiline(draw, "• " + line, x1 + 20, yy, 19, font(18), fill=(40, 40, 40), line_gap=4)


def arrow(draw, start, end, color="#566D63", width=4):
    draw.line((start, end), fill=color, width=width)
    import math

    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 15
    spread = 0.55
    p1 = (x2 - length * math.cos(angle - spread), y2 - length * math.sin(angle - spread))
    p2 = (x2 - length * math.cos(angle + spread), y2 - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill=color)


def create_architecture_image(path):
    w, h = 1800, 1120
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((w // 2, 45), "系统总体架构图", font=font(36, True), fill="#263B34", anchor="mm")
    draw.text((w // 2, 86), "微信小程序端 + Web 管理后台 + 服务器端 + 模拟客户端 + 数据/异步处理层", font=font(21), fill="#60746B", anchor="mm")

    draw_box(draw, (90, 140, 440, 310), "微信小程序端", ["登录与设备绑定", "播放控制与听歌数据", "好友一起听与分享"])
    draw_box(draw, (525, 140, 875, 310), "Web 管理后台", ["角色登录与鉴权", "运营统计与用户画像", "设备运维与反馈处理"])
    draw_box(draw, (960, 140, 1310, 310), "模拟客户端", ["设备上报模拟", "接口冒烟与压测", "无硬件环境验证"])
    draw_box(draw, (1395, 140, 1710, 310), "真实/模拟设备", ["音量/信号/低音上报", "播放状态上报", "连接与运行状态"])
    draw_box(draw, (270, 405, 680, 560), "Nginx / HTTPS 入口", ["统一公网域名", "反向代理", "跨域与访问入口"])
    draw_box(draw, (850, 405, 1260, 560), "Gunicorn + Flask 服务", ["RESTful API", "业务处理与权限控制", "健康检查与指标"])
    draw_box(draw, (90, 660, 390, 840), "小程序业务接口", ["/api/auth/*", "/api/device/*", "/api/player/*", "/api/listening-data/*"])
    draw_box(draw, (450, 660, 750, 840), "后台管理接口", ["/api/admin/*", "角色权限控制", "设备/反馈/统计"])
    draw_box(draw, (810, 660, 1110, 840), "数据库维护接口", ["/api/db/*", "表数据维护", "统计任务触发"])
    draw_box(draw, (1170, 660, 1470, 840), "设备上报接口", ["/api/bass", "/api/signal", "/api/volume", "/api/runtime"])
    draw_box(draw, (100, 930, 390, 1070), "MySQL", ["用户/设备/绑定", "播放历史/反馈", "运营统计日报"])
    draw_box(draw, (465, 930, 755, 1070), "MongoDB", ["设备运行状态", "播放器状态", "播放日志/队列"])
    draw_box(draw, (830, 930, 1120, 1070), "RabbitMQ", ["高频上报削峰", "异步消息缓冲"])
    draw_box(draw, (1195, 930, 1485, 1070), "Worker/统计任务", ["异步写入", "日志处理", "daily_stats_job"])

    for sx in [265, 700, 1135, 1550]:
        arrow(draw, (sx, 310), (480, 405))
    arrow(draw, (680, 482), (850, 482))
    for tx in [240, 600, 960, 1320]:
        arrow(draw, (1055, 560), (tx, 660))
    for sx, tx in [(240, 245), (600, 245), (960, 610), (1320, 975), (1320, 1340)]:
        arrow(draw, (sx, 840), (tx, 930))
    arrow(draw, (975, 930), (975, 840))
    arrow(draw, (1340, 930), (1040, 840))
    img.save(path, quality=95)


def er_box(draw, xy, title, fields, fill="#F8FBFA"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline="#4F6F64", width=3)
    draw.rectangle((x1, y1, x2, y1 + 38), fill="#4F6F64")
    draw.text(((x1 + x2) // 2, y1 + 19), title, font=font(19, True), fill="white", anchor="mm")
    yy = y1 + 48
    for field in fields[:6]:
        draw.text((x1 + 14, yy), field, font=font(15), fill="#263B34")
        yy += 24


def create_er_image(path):
    w, h = 1800, 1200
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((w // 2, 45), "核心 E-R 关系图", font=font(36, True), fill="#263B34", anchor="mm")
    draw.text((w // 2, 86), "围绕用户、设备、播放、音乐服务、社交、反馈和后台权限的数据关系", font=font(21), fill="#60746B", anchor="mm")
    boxes = {
        "user": (690, 145, 1020, 305),
        "profile": (300, 145, 610, 285),
        "token": (1100, 145, 1410, 285),
        "binding": (690, 405, 1020, 575),
        "device": (300, 405, 610, 575),
        "settings": (70, 650, 380, 790),
        "history": (690, 675, 1020, 850),
        "media": (1100, 675, 1410, 835),
        "music": (1450, 405, 1740, 575),
        "room": (300, 900, 610, 1070),
        "share": (690, 930, 1020, 1070),
        "feedback": (1100, 930, 1410, 1070),
        "admin": (1450, 910, 1740, 1070),
    }
    er_box(draw, boxes["user"], "`user` 用户", ["PK user_id", "username", "phone/email", "nickname/avatar", "status", "last_login_at"])
    er_box(draw, boxes["profile"], "user_profile", ["PK profile_id", "FK user_id", "age/gender", "province/city", "active_level", "bound_platforms"])
    er_box(draw, boxes["token"], "auth_token", ["PK auth_id", "FK user_id", "platform_type", "token/session", "expires_at"])
    er_box(draw, boxes["binding"], "user_device_binding", ["PK binding_id", "FK user_id", "FK device_id", "custom_device_name", "is_primary", "bind_time"])
    er_box(draw, boxes["device"], "device 设备", ["PK device_id", "device_number", "model_name", "online_status", "firmware_version", "last_active"])
    er_box(draw, boxes["settings"], "device_settings", ["PK setting_id", "FK device_id", "volume_limit", "night_mode", "power_save"])
    er_box(draw, boxes["history"], "play_history", ["PK history_id", "FK user_id", "FK device_id", "FK mapping_id", "play_duration", "created_at"])
    er_box(draw, boxes["media"], "media_mapping", ["PK mapping_id", "FK user_id", "song_title", "artist", "platform", "external_id"])
    er_box(draw, boxes["music"], "music_service_*", ["binding_id", "FK user_id", "service", "account_name", "permissions", "sync_status"])
    er_box(draw, boxes["room"], "listen_room_*", ["room_id", "owner_user_id", "member.user_id", "comment.user_id", "status"])
    er_box(draw, boxes["share"], "share_record", ["PK share_id", "FK user_id", "song_id", "room_id", "share_type", "created_at"])
    er_box(draw, boxes["feedback"], "user_feedback", ["PK feedback_id", "FK user_id", "feedback_type", "status", "priority", "admin_id"])
    er_box(draw, boxes["admin"], "admin_user/role", ["admin_id", "role_id", "permission_id", "login_log", "operation_log"])

    def center(key):
        x1, y1, x2, y2 = boxes[key]
        return ((x1 + x2) // 2, (y1 + y2) // 2)

    for a, b, label in [
        ("user", "profile", "1:1"),
        ("user", "token", "1:N"),
        ("user", "binding", "1:N"),
        ("device", "binding", "1:N"),
        ("device", "settings", "1:1"),
        ("binding", "history", "关联播放"),
        ("user", "history", "1:N"),
        ("history", "media", "N:1"),
        ("user", "music", "1:N"),
        ("user", "room", "创建/加入"),
        ("user", "share", "1:N"),
        ("user", "feedback", "1:N"),
        ("admin", "feedback", "处理"),
    ]:
        ax, ay = center(a)
        bx, by = center(b)
        arrow(draw, (ax, ay), (bx, by), color="#789186", width=3)
        mx, my = (ax + bx) // 2, (ay + by) // 2
        draw.rounded_rectangle((mx - 42, my - 16, mx + 42, my + 16), radius=8, fill="#EEF5F2", outline="#C7D7D1")
        draw.text((mx, my), label, font=font(14), fill="#263B34", anchor="mm")
    img.save(path, quality=95)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[name].font.name = "微软雅黑"
        doc.styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(11.5)
    doc.styles["Heading 3"].font.bold = True


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9.3)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(20)
    run.bold = True


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(96, 96, 96)


def para(doc, text, first=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    if first:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("• " + item)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10.5)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], header, bold=True, color=(255, 255, 255))
        set_cell_shading(t.rows[0].cells[i], "4F6F64")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_doc():
    create_architecture_image(ARCH_IMG)
    create_er_image(ER_IMG)

    doc = Document()
    style_doc(doc)
    add_title(doc, "智能音箱软件项目六部分报告")
    add_subtitle(doc, "按《改进方案.docx》要求重新整理：软件概述、需求分析、系统设计、系统实现、系统测试、部署与运维")
    add_subtitle(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    doc.add_heading("1. 软件概述", level=1)
    para(doc, "本系统实现的是一套面向智能音箱场景的多端音乐播放、设备管理与运营分析软件系统。系统不只是单一的音乐播放页面，也不是单纯的后台管理网站，而是将普通用户的小程序使用、智能音箱设备状态采集、服务器端数据处理、Web 后台运营管理和模拟客户端测试验证结合在一起的综合型软件系统。")
    para(doc, "从软件类型上看，本系统兼具智能硬件配套应用、B/S 架构 Web 管理系统、RESTful API 服务系统和数据统计分析系统的特点。微信小程序端负责用户日常操作，Web 管理后台负责运营和运维管理，服务器端负责接口、鉴权、数据存储和异步处理，模拟客户端负责设备数据模拟、接口联调和压力测试。")
    doc.add_heading("1.1 系统面向的用户", level=2)
    table(doc, ["用户类型", "使用入口", "主要需求"], [
        ["普通用户/智能音箱使用者", "微信小程序端", "登录系统、绑定设备、控制播放、查看听歌历史和周报、绑定音乐服务、参与好友一起听和分享。"],
        ["超级管理员", "Web 管理后台", "查看系统整体数据，管理后台账号和角色权限，掌握用户、设备、销售、活跃度和区域分布情况。"],
        ["市场分析管理员", "Web 管理后台", "查看热歌排行、用户画像、音乐服务分布、留存和用户价值，为运营策略提供数据支持。"],
        ["设备运维/普通管理员", "Web 管理后台", "查看设备列表、运行状态和日志，处理用户反馈，执行设备改名、解绑等运维操作。"],
        ["系统维护人员", "服务器、数据库、日志、容器平台", "完成部署、监控、备份、故障排查、安全更新和版本迭代。"],
        ["开发测试人员", "模拟客户端、接口脚本、JMeter、Postman 等", "模拟用户请求和设备上报，验证登录、查询、提交数据、接口性能和异常场景。"],
    ], widths=[3.2, 3.8, 8.2])
    doc.add_heading("1.2 系统解决的问题", level=2)
    bullets(doc, [
        "解决普通用户缺少统一智能音箱控制入口的问题，使用户可以通过小程序完成设备绑定、音乐播放、音量控制、音乐服务授权和听歌数据查看。",
        "解决设备状态、播放记录、用户行为和运营数据分散的问题，通过服务器端和数据库统一沉淀数据。",
        "解决 Web 后台无法及时了解小程序端真实使用情况的问题，使小程序端产生的播放、绑定、分享、反馈等数据可以进入后台统计和管理流程。",
        "解决后台管理动作无法反馈到用户侧的问题，使 Web 端修改设备名称、解绑设备、处理反馈等操作能够通过共享数据源被小程序端获取。",
        "解决高频设备上报直接写数据库导致压力过大的问题，通过 RabbitMQ、Worker、短期缓存和定时统计任务实现削峰和汇总。",
    ])
    doc.add_heading("1.3 软件主要目标和预期效果", level=2)
    table(doc, ["目标", "说明", "预期效果"], [
        ["用户体验目标", "提供微信小程序端，覆盖登录、设备、播放、数据、好友和授权流程。", "普通用户能够用较低学习成本完成智能音箱日常使用。"],
        ["数据联通目标", "以服务器端和数据库作为统一数据中心，避免小程序和 Web 各自维护孤立数据。", "小程序产生的数据能被 Web 后台观察，Web 后台修改也能反馈到小程序。"],
        ["管理分析目标", "为不同管理员角色提供数据总览、市场分析、用户画像、设备运维和反馈处理。", "管理人员可以基于真实数据进行运营决策和系统维护。"],
        ["性能稳定目标", "通过 Gunicorn、Nginx、RabbitMQ、Worker、短 TTL 缓存和每日统计任务提升并发能力。", "在设备频繁上报和后台频繁查询时，系统仍能保持稳定响应。"],
        ["可扩展目标", "前后端分离，业务模块拆分，数据库表按用户、设备、播放、反馈、统计等领域设计。", "后续可扩展更多音乐平台、更多设备类型、更多后台页面和更多客户端入口。"],
    ], widths=[3.0, 6.4, 6.0])
    doc.add_heading("1.4 多端系统组成", level=2)
    table(doc, ["端/组件", "项目位置或技术", "作用"], [
        ["微信小程序端", "wxapp/miniprogram-1/miniprogram，TypeScript", "普通用户入口，包含登录、首页、好友、数据、设备、历史、音乐授权等页面。"],
        ["Web 管理后台", "src/App.vue、src/api.js，Vue3 + Vite + Element Plus", "管理员入口，提供数据总览、统计分析、设备运维、反馈处理、账号权限管理。"],
        ["服务器端", "app.py、api_pkg/*、admin_routes.py，Flask", "提供 RESTful API、用户身份识别、管理员鉴权、业务逻辑、CORS、健康检查和统计触发。"],
        ["数据与异步处理端", "MySQL、MongoDB、RabbitMQ、worker_*.py、daily_stats_job.py", "保存业务数据和运行状态，处理高频上报、异步写入、日志记录和每日统计。"],
        ["模拟客户端/测试端", "music/simulator.py、project_player、test_all_api_mock.py、scripts/smoke_existing_apis.py、JMeter 文件", "模拟设备上报、播放事件、接口调用和并发压力，辅助联调和验收。"],
    ], widths=[3.2, 5.2, 7.0])

    doc.add_heading("2. 需求分析", level=1)
    para(doc, "需求分析不能只停留在功能列表层面，而应从不同用户的实际诉求、系统功能模块和非功能质量要求三个层次展开。本系统同时存在普通用户端和管理端，因此需要明确小程序和 Web 后台的功能差异：小程序侧更关注个人使用和设备交互，Web 侧更关注管理、统计、运维和权限。")
    doc.add_heading("2.1 用户需求分析", level=2)
    table(doc, ["用户", "需求描述", "对应系统能力"], [
        ["普通用户", "需要登录系统、绑定智能音箱、查看设备状态、控制播放、绑定音乐服务、查询听歌历史和周报、提交反馈。", "微信登录、设备接口、播放器接口、音乐服务接口、听歌数据接口、反馈接口。"],
        ["管理员", "需要管理后台账号、查看统计数据、维护用户和设备、处理反馈、分析市场运营情况。", "后台登录鉴权、角色权限、数据总览、用户画像、热歌排行、设备运维、反馈处理。"],
        ["系统维护人员", "需要部署服务、监控运行状态、备份数据库、查看日志、处理异常和进行版本迭代。", "Docker Compose、Nginx、Gunicorn、/health、/internal/metrics、日志和数据库备份。"],
        ["开发测试人员", "需要在无真实硬件或部分硬件不可用时模拟请求和设备数据，验证接口和性能。", "模拟客户端、接口冒烟脚本、JMeter 压测文件、测试数据生成脚本。"],
    ], widths=[3.2, 6.4, 6.2])
    doc.add_heading("2.2 系统功能需求", level=2)
    table(doc, ["功能模块", "主要功能", "涉及端/接口"], [
        ["用户登录注册模块", "小程序微信登录、后台管理员登录、验证码/短信验证码、token 生成与校验。", "/api/auth/wechat-login、/api/admin/login、/api/admin/profile"],
        ["设备管理模块", "设备搜索、绑定、详情、电量、省电模式、电量提醒、改名、解绑、运行状态查看。", "/api/device/*、/api/admin/operator/device/*"],
        ["播放控制模块", "播放歌曲、音量调节、播放/暂停/切歌控制、加入下一首、播放状态记录。", "/api/player/play-song、/api/player/volume、/api/player/control、/api/player/add-next"],
        ["数据管理模块", "用户资料、播放历史、音乐服务绑定、分享记录、反馈、设备日志等数据维护。", "MySQL 表、Mongo 集合、/api/db/*"],
        ["查询统计模块", "听歌概览、周报、热歌排行、用户画像、活跃率、留存、区域热力图、设备统计。", "/api/listening-data/*、/api/admin/super/*、/api/admin/market/*"],
        ["后台管理模块", "管理员账号、角色权限、用户反馈、设备运维、公告、导出和后台个人信息维护。", "admin_routes.py、src/App.vue"],
        ["接口通信模块", "小程序与 Web 均通过 RESTful API 与 Flask 后端通信，使用 JSON 数据格式。", "wx.request、fetch、Authorization 请求头"],
        ["权限管理模块", "后台按 super_admin、market_admin、operator_admin 分权；小程序侧按用户 token 隔离数据。", "管理员 token、用户 token、角色权限表"],
        ["日志管理模块", "设备日志、后台操作日志、登录日志、Worker 日志和网关指标。", "device_log、admin_operation_log、admin_login_log、/internal/metrics"],
    ], widths=[3.1, 6.4, 6.1])
    doc.add_heading("2.3 非功能需求", level=2)
    table(doc, ["需求类型", "具体要求", "本项目对应设计"], [
        ["性能需求", "常规页面请求应保持较快响应；设备高频上报不能直接压垮数据库；后台统计可接受短时间缓存。", "Gunicorn 多进程多线程、RabbitMQ 削峰、Worker 异步写入、后台短 TTL 缓存。"],
        ["安全需求", "用户和管理员需要鉴权；密码和 token 不能明文泄露；后台不同角色访问不同接口。", "Authorization: Bearer token、管理员角色权限、密码哈希、ADMIN_TOKEN_SECRET、401/403 权限控制。"],
        ["可用性需求", "界面操作简单，错误提示明确，小程序端适配移动端，后台适合管理人员高频查看。", "小程序页面分为首页/好友/数据/设备；Web 使用 Element Plus 表格、弹窗、提示和导出。"],
        ["可维护性需求", "系统结构应分层清晰，接口模块化，方便定位问题和扩展功能。", "Flask Blueprint 拆分 auth/device/player/music/admin/db 等模块，前端统一 request 封装。"],
        ["可扩展性需求", "后续可接入更多音乐平台、更多设备类型、更多运营指标和更多客户端。", "music_service_binding 按 service 扩展，设备表支持 model/device_type，接口采用 RESTful JSON。"],
        ["可靠性需求", "数据库异常、接口失败、设备离线等场景需要可控处理，并应有数据备份和日志。", "健康检查、日志表、运行指标、兜底数据/空状态、数据库备份策略。"],
        ["兼容性需求", "Web 端应兼容主流浏览器，小程序端应适配微信开发者工具和真机环境。", "Vue3/Vite 浏览器访问，小程序 appid 与 miniprogramRoot 配置，微信 request 合法域名。"],
    ], widths=[3.0, 6.2, 6.4])

    doc.add_heading("3. 系统设计", level=1)
    para(doc, "系统设计部分按照改进要求补充架构图、技术栈、模块设计、数据库设计、E-R 图、接口设计和表关系说明。整体采用前后端分离与分层架构：小程序端和 Web 端只负责交互展示，后端负责业务处理和鉴权，数据库负责持久化，消息队列和 Worker 负责异步处理。")
    doc.add_heading("3.1 架构设计", level=2)
    doc.add_picture(str(ARCH_IMG), width=Cm(16.4))
    caption(doc, "图 3-1 系统总体架构图")
    para(doc, "系统采用“小程序端 + Web 管理后台 + Flask 后端服务 + MySQL/MongoDB + RabbitMQ/Worker”的多端架构。客户端通过 RESTful API 与服务器通信，接口数据格式为 JSON。Web 后台请求会携带管理员 Authorization token；小程序端后续应携带用户 token，确保用户数据隔离。Nginx 作为公网入口，Gunicorn 承载 Flask 服务，RabbitMQ 和 Worker 处理高频设备上报与异步写入。")
    doc.add_heading("3.2 技术栈设计", level=2)
    table(doc, ["层级", "技术/组件", "说明"], [
        ["微信小程序端", "TypeScript、微信小程序框架、wx.request", "实现移动端用户页面和后端 API 通信。"],
        ["Web 管理后台", "Vue3、Vite、Element Plus、fetch", "实现后台单页应用、数据表格、弹窗、表单、图表和权限页面。"],
        ["服务器端", "Python、Flask、Blueprint、Gunicorn", "实现 RESTful API、业务逻辑、权限校验、CORS、健康检查。"],
        ["关系数据库", "MySQL", "保存用户、设备、绑定关系、播放历史、反馈、权限和统计日报。"],
        ["文档/状态数据库", "MongoDB", "保存设备运行状态、播放器状态、播放日志、播放队列等变化频繁数据。"],
        ["消息队列", "RabbitMQ", "承接高频设备上报，实现削峰和异步处理。"],
        ["部署运维", "Docker Compose、Nginx、Portainer、/internal/metrics", "实现容器编排、反向代理、运行监控和服务管理。"],
    ], widths=[3.0, 4.5, 8.0])
    doc.add_heading("3.3 模块设计", level=2)
    table(doc, ["模块", "主要职责", "对应文件/接口"], [
        ["认证与权限模块", "处理微信用户登录、后台管理员登录、token 校验、角色权限控制。", "api_pkg/auth_home.py、auth_routes.py、admin_routes.py、security_utils.py"],
        ["设备模块", "处理设备绑定、详情、电量、设置、改名、解绑和运行状态。", "api_pkg/device.py、device_routes.py、/api/admin/operator/device/*"],
        ["播放模块", "处理播放、音量、控制、下一首、播放历史和播放器状态。", "api_pkg/player.py、Mongo player_state/play_logs、MySQL play_history"],
        ["音乐服务模块", "处理 QQ 音乐、网易云等服务绑定、权限、同步状态。", "api_pkg/music_service.py、music_service_binding、music_service_permission"],
        ["社交与分享模块", "处理好友、一起听房间、评论、分享链接和分享卡片。", "api_pkg/friends_rooms.py、api_pkg/share_data_history.py"],
        ["后台统计模块", "处理用户数、设备数、热歌、用户画像、留存、区域热力图和反馈列表。", "admin_routes.py、daily_stats_job.py、daily_stats 等统计表"],
        ["数据库维护模块", "提供数据库查询、详情、增删改和统计任务触发能力。", "db_api_service.py、api_pkg/db_board.py、/api/db/*"],
        ["异步处理模块", "处理设备上报、消息队列、日志、校验、写入和自动扩缩容。", "mq_config.py、worker_*.py、auto_scaler.py"],
    ], widths=[3.0, 6.0, 6.4])
    doc.add_heading("3.4 数据库设计与 E-R 图", level=2)
    doc.add_picture(str(ER_IMG), width=Cm(16.4))
    caption(doc, "图 3-2 核心 E-R 关系图")
    para(doc, "数据库设计以用户、设备和播放行为为核心。用户表与用户画像、登录 token、设备绑定、播放历史、音乐服务绑定、分享记录和反馈记录形成一对一或一对多关系；设备表与设备设置、绑定关系、运行状态和日志关联；后台管理员表、角色表、权限表用于 Web 后台的分权管理。")
    table(doc, ["表名", "主键", "核心字段", "关系/说明"], [
        ["user", "user_id", "username、phone、nickname、avatar、status、last_login_at", "普通用户主表，与 user_profile、auth_token、user_device_binding、play_history 等关联。"],
        ["user_profile", "profile_id", "user_id、age、province_name、active_level、value_level、bound_platforms", "用户画像表，与 user 一对一。"],
        ["auth_token", "auth_id", "user_id、platform_type、token、expires_at", "保存用户登录凭证，与 user 多对一。"],
        ["device", "device_id", "device_number、model_name、status、online_status、firmware_version", "设备主表，通过 user_device_binding 与用户关联。"],
        ["user_device_binding", "binding_id", "user_id、device_id、custom_device_name、is_primary、bind_time", "用户和设备的绑定关系表，支持一人多设备和设备归属。"],
        ["play_history", "history_id", "user_id、device_id、mapping_id、play_duration、created_at", "播放历史表，是听歌统计、热歌排行和用户活跃分析的基础。"],
        ["media_mapping", "mapping_id", "user_id、song_title、artist、platform、external_id、cover_url", "统一不同音乐平台歌曲信息。"],
        ["music_service_binding", "binding_id", "user_id、service、account_name、sync_status、bound_at", "记录第三方音乐服务绑定状态。"],
        ["listen_room / member / comment", "room_id / id / comment_id", "owner_user_id、room_code、member.user_id、content、status", "一起听房间、成员和评论表。"],
        ["share_record", "share_id", "user_id、song_id、room_id、share_type、share_url、image_url", "记录用户分享行为。"],
        ["user_feedback", "feedback_id", "user_id、feedback_type、title、status、priority、admin_id", "记录用户反馈和后台处理状态。"],
        ["admin_user / admin_role / admin_permission", "admin_id / role_id / permission_id", "username、role_code、permission_code", "后台账号、角色和权限控制。"],
        ["daily_stats 等统计表", "stat_date 或 id", "play_count、active_user_count、online_device_count、sales_amount", "后台总览、趋势、留存、热歌和用户价值统计。"],
    ], widths=[3.1, 2.4, 5.2, 5.2])
    doc.add_heading("3.5 接口设计", level=2)
    table(doc, ["接口", "方法", "使用端", "功能说明", "主要数据源"], [
        ["/api/auth/wechat-login", "POST", "小程序", "微信登录，建立用户身份和登录态。", "user、auth_token"],
        ["/api/home/overview", "GET", "小程序", "首页概览，返回设备、播放状态和历史统计。", "device、user_device_binding、play_history、Mongo player_state"],
        ["/api/device/list / detail / bind / unbind", "GET/POST", "小程序", "设备列表、详情、绑定和解绑。", "device、user_device_binding、device_bind_task"],
        ["/api/player/play-song", "POST", "小程序", "播放歌曲并写入播放历史和播放器状态。", "play_history、media_mapping、Mongo player_state/play_logs"],
        ["/api/player/volume / control", "POST", "小程序", "调整音量、播放暂停等控制。", "Mongo device_runtime、player_state"],
        ["/api/music-service/*", "GET/POST", "小程序", "音乐服务绑定、授权、同步进度和解绑。", "music_service_binding、music_service_permission"],
        ["/api/listening-data/*", "GET/POST", "小程序", "听歌概览、周报、生成分享卡片。", "play_history、daily_stats、user_activity_daily"],
        ["/api/friends/* / listen-room/*", "GET/POST", "小程序", "好友、一起听房间、成员和评论。", "friendship、listen_room、listen_room_member、listen_room_comment"],
        ["/api/share/song-link / song-card", "POST", "小程序", "生成歌曲分享链接或卡片。", "share_record、media_mapping"],
        ["/api/admin/login / profile", "GET/POST", "Web 后台", "管理员登录、个人信息和权限识别。", "admin_user、admin_role、admin_permission"],
        ["/api/admin/super/* / market/*", "GET", "Web 后台", "总览统计、市场分析、热歌、画像、留存。", "daily_stats、play_history、user_profile、sales_order"],
        ["/api/admin/operator/device/*", "GET/POST", "Web 后台", "设备列表、状态、日志、改名和解绑。", "device、user_device_binding、device_log、Mongo device_runtime"],
    ], widths=[4.1, 1.7, 2.3, 5.0, 4.0])

    doc.add_heading("4. 系统实现", level=1)
    para(doc, "系统实现部分按照多端系统要求分别说明模拟客户端、服务器端、微信小程序端和 Web 管理后台的实现方式。当前项目已经具备四类实现基础：模拟端用于数据生成与压测，服务器端用于接口和数据处理，小程序端用于用户操作，Web 端用于后台管理。")
    doc.add_heading("4.1 模拟客户端实现", level=2)
    para(doc, "模拟客户端主要用于在没有完整智能音箱硬件或需要进行压力测试时生成测试数据。项目中的 music/simulator.py 会模拟设备数据并写入 CSV 日志，project_player/data_gen.py 可生成设备信号、音量、低音、点赞等数据包，test_all_api_mock.py 和 scripts/smoke_existing_apis.py 用于模拟接口请求，jmeter_gateway_stress_test.jmx 用于压力测试。")
    bullets(doc, ["模拟设备上报：生成设备 ID、信号强度、音量、低音、点赞、时间戳等数据。", "模拟用户请求：调用登录、设备、播放、听歌数据、分享、一起听等接口。", "模拟压力场景：通过 JMeter 或脚本向网关接口发送并发请求，观察响应时间和失败率。", "保存验证材料：CSV 日志、接口响应、后端日志和数据库写入结果可作为报告截图材料。"])
    doc.add_heading("4.2 服务器端实现", level=2)
    para(doc, "服务器端使用 Python Flask 实现，app.py 是主入口，注册多个 Blueprint。后端分为接口层、业务处理层、数据访问层和异步处理层。接口层接收小程序、Web 后台和设备上报请求；业务处理层完成用户身份、设备绑定、播放控制、音乐服务、反馈和统计逻辑；数据访问层读写 MySQL 与 MongoDB；异步处理层通过 RabbitMQ 和 Worker 消费高频上报数据。")
    table(doc, ["实现点", "说明", "对应文件"], [
        ["Controller/路由层", "使用 Flask route 和 Blueprint 暴露 RESTful API。", "app.py、api_pkg/*、admin_routes.py"],
        ["Service/业务逻辑", "处理登录、设备、播放、统计、反馈、权限等业务。", "api_pkg/common.py、daily_stats_job.py、admin_routes.py"],
        ["DAO/数据访问", "连接 MySQL、MongoDB，执行查询、写入、聚合和数据库维护。", "db_config.py、db_api_service.py、storage_backends.py"],
        ["权限控制", "后台 token 鉴权、角色权限判断；用户侧应按 token 解析 user_id。", "admin_routes.py、security_utils.py、api_pkg/common.py"],
        ["异常处理", "接口失败返回 JSON 错误，设备上报队列满可返回 503，数据库异常可进入兜底或空状态。", "app.py、api_pkg/common.py"],
        ["日志记录", "设备日志、后台操作日志、登录日志和 Worker 日志。", "device_log、admin_operation_log、worker_logger.py"],
    ], widths=[3.0, 6.6, 5.8])
    doc.add_heading("4.3 微信小程序实现", level=2)
    para(doc, "微信小程序端位于 wxapp/miniprogram-1/miniprogram，使用 TypeScript 编写。app.json 配置了 login、index、friends、data、device、history、auth、logs 等页面，并使用自定义 tabBar。services/api.ts 封装后端接口类型和请求方法，utils/auth.ts 保存本地登录态，pages/login/index.ts 调用 wx.login。")
    table(doc, ["页面/模块", "功能", "与后端交互"], [
        ["login", "用户登录、获取微信登录 code、保存本地登录态。", "/api/auth/wechat-login"],
        ["index", "首页概览，展示设备、电量、播放状态和历史数量。", "/api/home/overview"],
        ["device", "设备详情、电量、绑定、改名、省电、电量提醒和解绑。", "/api/device/*"],
        ["data", "听歌概览、周报、生成卡片。", "/api/listening-data/*"],
        ["friends", "好友正在听、搜索好友、一起听房间。", "/api/friends/*、/api/listen-room/*"],
        ["history", "播放历史列表、删除或清理历史。", "/api/play-history/*"],
        ["auth", "音乐服务绑定、权限设置、同步进度。", "/api/music-service/*"],
    ], widths=[2.6, 6.0, 6.8])
    para(doc, "需要注意的是，当前小程序请求封装中普通用户接口还需要进一步补齐 Authorization: Bearer <token>，否则后端无法稳定识别真实微信用户。正式联通时应将 wx.login 换取的后端 token 保存到登录态，并在所有用户侧请求中携带。")
    doc.add_heading("4.4 Web 管理后台实现", level=2)
    para(doc, "Web 管理后台位于 src 目录，使用 Vue3、Vite 和 Element Plus 实现。src/api.js 统一封装 API_BASE、请求参数、Authorization 请求头、JSON 解析和错误处理；src/App.vue 根据管理员角色加载不同后台功能，包括数据总览、市场分析、设备运维、反馈处理、账号管理、密码修改、公告创建和导出等。")
    bullets(doc, ["登录鉴权：管理员登录后保存 admin_token，后续请求通过 Authorization 请求头访问后台接口。", "数据管理：展示用户、设备、反馈、统计数据，支持设备重命名、解绑、反馈状态更新、账号管理等操作。", "图表统计：通过后台统计接口读取热歌排行、用户画像、趋势增长、区域热力图、留存和用户价值。", "权限管理：根据 super_admin、market_admin、operator_admin 控制可见菜单和可访问接口。", "接口交互：Web 端调用 /api/admin/* 和 /api/db/*，与小程序共享底层数据源。"])
    doc.add_heading("4.5 数据联通实现说明", level=2)
    para(doc, "小程序端和 Web 端的数据联通以“同一后端、同一数据库、同一用户和设备标识”为核心。小程序播放歌曲会写入 play_history 和 Mongo player_state/play_logs，Web 后台可读取热歌排行、趋势和设备状态；Web 后台修改 device 或 user_device_binding 后，小程序重新拉取 /api/device/detail 或 /api/home/overview 即可看到变化。对于日报、留存和画像等聚合指标，需要通过 daily_stats_job 定时或手动汇总。")

    doc.add_heading("5. 系统测试", level=1)
    para(doc, "系统测试不能只写“测试通过”，而应包含测试设计、测试用例、测试结果记录和截图/日志证据。本项目可从功能测试、接口测试、性能测试、兼容性测试和异常测试五个方面展开。当前项目已包含测试脚本和压测文件，但当前本机环境未安装 pytest，因此 pytest 单元测试需要先安装依赖后执行；接口冒烟脚本和 JMeter 文件可作为后续验收材料。")
    doc.add_heading("5.1 测试设计", level=2)
    table(doc, ["测试类型", "测试目标", "测试工具/材料"], [
        ["功能测试", "验证小程序登录、设备、播放、音乐服务、数据周报、好友一起听，以及 Web 后台登录、统计、设备运维、反馈处理。", "小程序真机/开发者工具、浏览器、手工测试用例。"],
        ["接口测试", "验证后端 API 状态码、响应字段、权限控制和数据库写入。", "Postman、test_all_api_mock.py、scripts/smoke_existing_apis.py。"],
        ["性能测试", "验证设备上报和后台查询在并发情况下的响应时间和失败率。", "jmeter_gateway_stress_test.jmx、/internal/metrics、RabbitMQ 管理界面。"],
        ["兼容性测试", "验证 Web 在主流浏览器、小程序在开发者工具和真机中的显示和请求。", "Chrome/Edge、微信开发者工具、微信真机预览。"],
        ["异常测试", "验证未登录、token 失效、跨角色访问、设备不存在、数据库异常、队列满等情况。", "接口脚本、后端日志、数据库和队列状态。"],
    ], widths=[3.0, 7.0, 5.6])
    doc.add_heading("5.2 测试用例", level=2)
    table(doc, ["编号", "测试项", "操作步骤", "预期结果"], [
        ["TC-01", "小程序登录", "打开小程序登录页，调用 wx.login 并请求 /api/auth/wechat-login。", "登录成功后获得用户身份，后续请求可按用户隔离数据。"],
        ["TC-02", "设备绑定", "搜索附近设备，选择设备并提交绑定。", "user_device_binding 产生绑定记录，小程序设备页显示真实设备。"],
        ["TC-03", "播放歌曲", "在小程序发起 /api/player/play-song。", "play_history 写入记录，Mongo player_state 更新，Web 后台热歌/状态可查询。"],
        ["TC-04", "Web 设备改名", "后台管理员登录后对设备执行 rename。", "device 或绑定名称更新，小程序刷新后显示新名称。"],
        ["TC-05", "Web 设备解绑", "后台对某用户设备执行 unbind。", "绑定关系解除，小程序刷新后显示无设备或引导绑定。"],
        ["TC-06", "音乐服务绑定", "小程序绑定 QQ/网易云等服务并设置权限。", "music_service_binding 和 permission 表有记录，后台音乐服务分布可统计。"],
        ["TC-07", "反馈处理", "小程序提交反馈，后台查看并更新状态。", "user_feedback 状态更新，process_log 记录处理过程。"],
        ["TC-08", "权限控制", "使用 market_admin 访问 super_admin 接口，或未登录访问后台 profile。", "跨角色返回 403，未登录返回 401。"],
        ["TC-09", "接口异常", "请求不存在设备、错误参数或空 token。", "接口返回明确错误码和提示，不产生脏数据。"],
        ["TC-10", "性能压测", "使用 JMeter 对设备上报接口发起并发请求。", "失败率在可接受范围，/internal/metrics 和队列状态正常。"],
    ], widths=[1.4, 2.8, 7.0, 5.0])
    doc.add_heading("5.3 测试结果记录方式", level=2)
    table(doc, ["结果类型", "应记录内容", "建议证据"], [
        ["功能测试结果", "每个功能是否通过、失败原因、修复状态。", "小程序页面截图、Web 页面截图、操作录屏。"],
        ["接口测试结果", "请求 URL、方法、参数、状态码、响应 JSON、数据库变化。", "Postman 截图、脚本输出、后端日志、数据库查询截图。"],
        ["性能测试结果", "并发数、总请求数、失败率、平均响应时间、P95/P99、吞吐量。", "JMeter 报告、/internal/metrics、RabbitMQ 队列截图。"],
        ["兼容性测试结果", "浏览器版本、小程序基础库版本、真机型号、显示异常记录。", "浏览器截图、微信开发者工具截图、真机截图。"],
        ["异常测试结果", "401/403/404/500 等错误场景是否符合预期，是否有日志和告警。", "接口响应截图、后台日志、错误处理说明。"],
    ], widths=[3.0, 6.4, 6.0])
    para(doc, "本次生成报告时，本地 Python 环境缺少 pytest，无法直接运行 tests/test_security_config.py 和 tests/test_mongo_event_logs.py。正式提交前建议执行：pip install pytest 后运行 py -m pytest tests -q，并将结果截图补充到报告附件或测试章节中。")

    doc.add_heading("6. 部署与运维", level=1)
    para(doc, "部署与运维部分需要写清楚环境、步骤和长期维护策略。本项目支持服务器部署，后端可通过 Nginx + Gunicorn + Flask 对外提供服务，Docker Compose 管理 RabbitMQ、gateway、worker、daily-stats-job 和 Portainer。Web 前端可通过 Vite 构建后部署，微信小程序通过微信开发者工具预览和上传。")
    doc.add_heading("6.1 部署环境", level=2)
    table(doc, ["类别", "环境/版本", "说明"], [
        ["操作系统", "Linux 服务器 / Windows 本地开发环境", "服务器用于生产部署，本地用于开发和文档生成。"],
        ["服务器", "8.137.165.220（现有部署说明中记录）", "作为后端公网访问入口或反向代理入口。"],
        ["后端环境", "Python、Flask 3.0.0、Gunicorn 22.0.0、PyMySQL、pymongo、pika", "requirements.txt 中记录主要依赖。"],
        ["Web 前端", "Node.js、Vue 3.5、Vite、Element Plus", "package.json 中记录依赖，npm run build 生成 dist。"],
        ["小程序端", "微信开发者工具、TypeScript、appid wxf9d56e7302242aff", "project.config.json 配置 miniprogramRoot=miniprogram/。"],
        ["数据库", "MySQL、MongoDB", "MySQL 保存业务表，MongoDB 保存运行状态和日志。"],
        ["消息队列", "RabbitMQ 3-management", "用于设备上报削峰和 Worker 异步处理。"],
        ["反向代理", "Nginx", "负责 HTTP/HTTPS 入口、反向代理到 Flask/Gunicorn。"],
        ["端口", "后端 5000，RabbitMQ 5672/15672，Portainer 9000，Nginx 80/443", "具体端口以服务器配置和安全组开放情况为准。"],
    ], widths=[3.0, 5.0, 7.0])
    doc.add_heading("6.2 部署步骤", level=2)
    bullets(doc, [
        "克隆或同步 Git 仓库到服务器项目目录，例如 /www/wwwroot/mysite。",
        "安装后端依赖：pip install -r requirements.txt，或使用 Dockerfile.gateway / Dockerfile.worker 构建镜像。",
        "初始化数据库：执行 init_smart_speaker_web.sql 和 sql/daily_stats.sql，准备 MySQL 表结构；配置 MongoDB 数据库。",
        "配置环境变量：MYSQL_HOST、MYSQL_DATABASE、MYSQL_USER、MYSQL_PASSWORD、MONGO_URI、MQ_USER、MQ_PASSWORD、ADMIN_TOKEN_SECRET、管理员密码等。",
        "启动基础服务：docker compose up -d --build rabbitmq gateway worker-writer worker-logger worker-validator worker-reader。",
        "按需要启动统计任务：docker compose --profile jobs up -d --build daily-stats-job。",
        "配置 Nginx，将公网域名或 IP 的请求反向代理到 127.0.0.1:5000，并配置 HTTPS、CORS 和超时。",
        "构建 Web 前端：npm install 后执行 npm run build，将 dist 部署到 Web 访问目录或通过 Nginx 提供静态资源。",
        "使用微信开发者工具打开 wxapp/miniprogram-1，确认 appid、request 合法域名和 miniprogramRoot 配置正确，进行预览和上传。",
        "部署后验证：访问 /health、/api/_health、/api/admin/captcha、后台登录、小程序首页、设备接口和 /internal/metrics。",
    ])
    code(doc, "docker compose up -d --build rabbitmq gateway worker-writer worker-logger worker-validator worker-reader")
    code(doc, "docker compose --profile jobs up -d --build daily-stats-job")
    code(doc, "npm run build")
    doc.add_heading("6.3 运维策略", level=2)
    table(doc, ["运维项", "策略"], [
        ["数据库备份", "定期备份 MySQL 和 MongoDB；重要版本发布前单独备份业务表和配置文件。"],
        ["日志管理", "按日期归档 Nginx、Flask、Worker、设备日志、后台操作日志，避免日志无限增长。"],
        ["异常监控", "通过 /health、/internal/metrics、RabbitMQ 队列长度、容器状态和错误日志判断服务是否异常。"],
        ["服务器重启策略", "使用 Docker restart policy 和 Gunicorn 平滑重载；发布时优先最小范围重启，避免影响数据库和队列。"],
        ["安全更新", "定期更新依赖和系统补丁；管理员密码、ADMIN_TOKEN_SECRET、数据库密码不得提交到 Git。"],
        ["权限管理", "后台账号按角色最小权限分配；高风险操作如解绑设备、删除账号、修改权限应记录操作日志。"],
        ["版本迭代", "通过 Git 分支和发布记录管理迭代；每次发布前执行接口冒烟、数据库备份和回滚方案确认。"],
        ["生产数据保护", "正式环境建议关闭演示兜底数据和 DAILY_STATS_AUTO_DEMO_DATA，避免模拟数据污染真实统计。"],
    ], widths=[3.2, 12.0])
    para(doc, "总体来看，本项目已经具备多端系统的部署基础。后续若要进一步提升生产可用性，应重点完善小程序用户 token、统一 Web 与小程序域名、关闭生产演示兜底、补充自动化测试流水线、增加监控告警和数据库定期备份机制。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("智能音箱软件项目六部分报告 - 按改进方案重生成")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.core_properties.title = "智能音箱软件项目六部分报告-按改进方案重生成"
    doc.core_properties.subject = "根据改进方案重写的软件概述、需求分析、系统设计、系统实现、系统测试、部署与运维"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build_doc()
    print(out)
    print(out.stat().st_size)
    print(ARCH_IMG)
    print(ER_IMG)
