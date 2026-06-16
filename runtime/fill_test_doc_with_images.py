from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


FONT = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_cell(cell):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, size=10.5, bold=False, align=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def picture_width(path: Path):
    with Image.open(path) as im:
        w, h = im.size
    if h > w * 1.35:
        return Cm(4.2)
    if w > h * 1.7:
        return Cm(9.2)
    return Cm(7.4)


def add_evidence(cell, intro, entries):
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(intro)
    set_run_font(run, size=9.5)

    for caption, image_path in entries:
        cap = cell.add_paragraph()
        cap.paragraph_format.space_before = Pt(4)
        cap_run = cap.add_run(caption)
        set_run_font(cap_run, size=9, bold=True)

        pic_p = cell.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic_p.add_run()
        pic_run.add_picture(str(image_path), width=picture_width(image_path))


def set_table_widths(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(3.0), Cm(7.0), Cm(17.2)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def main():
    desktop = Path.home() / "Desktop"
    src = desktop / "测试.docx"
    out = desktop / "测试_补齐截图版.docx"
    assets = Path.cwd() / "runtime" / "test_doc_assets" / "selected"

    doc = Document(str(src))
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(10.5)

    if doc.paragraphs:
        title = doc.paragraphs[0]
        title.text = ""
    else:
        title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("测试结果证据表（已补充截图）")
    set_run_font(title_run, size=16, bold=True)

    table = doc.tables[0]
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    set_table_widths(table)

    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=10.5, bold=True)

    row_data = {
        "功能测试结果": {
            "content": "已覆盖小程序首页、设备管理、Web 管理端数据总览等主要功能页面的展示证据；后续联调时可继续补充完整操作录屏。",
            "intro": "已补充小程序端与 Web 管理端主要页面截图，作为功能测试通过情况的界面证据。",
            "entries": [
                ("证据 1：小程序首页功能展示（来源：小程序设计草图.pdf 第 2 页）", assets / "function_miniprogram_home.jpg"),
                ("证据 2：Web 管理端数据总览页面（来源：web端设计草图.pdf 第 14 页）", assets / "function_web_dashboard.jpg"),
            ],
        },
        "接口测试结果": {
            "content": "API 功能测试返回 status=success、状态码 200；同时保留 RabbitMQ 监控页面，用于证明请求链路与消息队列处理正常。",
            "intro": "已补充接口请求成功返回、消息队列链路监控截图，覆盖 URL 请求、状态码、响应 JSON 与后端链路证据。",
            "entries": [
                ("证据 1：API 功能测试成功返回截图（来源：第三阶段.pptx 第 7 页）", assets / "interface_api_success.jpg"),
                ("证据 2：RabbitMQ 全链路监控截图（来源：第三阶段.pptx 第 7 页）", assets / "interface_rabbitmq_monitor.jpg"),
            ],
        },
        "性能测试结果": {
            "content": "已记录 JMeter 压力测试报告和 MQ 流量截图。资源中显示业务数约 84.26/s、最大请求数约 158.12/s，可作为性能测试结果依据。",
            "intro": "已补充 JMeter 报告与消息队列压力流量截图，覆盖并发测试、吞吐量和队列处理能力相关证据。",
            "entries": [
                ("证据 1：JMeter 压力测试报告截图（来源：第三阶段.pptx 第 14 页）", assets / "performance_jmeter_report.jpg"),
                ("证据 2：消息队列压力流量截图（来源：第三阶段.pptx 第 11 页）", assets / "performance_mq_traffic.jpg"),
            ],
        },
        "兼容性测试结果": {
            "content": "已补充小程序端与 Web 端页面截图，用于说明移动端和浏览器端界面展示情况；正式验收时建议继续记录浏览器版本、微信基础库版本和真机型号。",
            "intro": "已补充移动端小程序页面与 Web 登录入口页面，用于支撑兼容性测试中的多端界面检查。",
            "entries": [
                ("证据 1：小程序设备管理页面（来源：小程序设计草图.pdf 第 5 页）", assets / "compat_miniprogram_device.jpg"),
                ("证据 2：Web 登录/角色入口页面（来源：web端设计草图.pdf 第 17 页）", assets / "compat_web_login.jpg"),
            ],
        },
        "异常测试结果": {
            "content": "已补充 Token 错误场景证据：请求被服务器拒绝并返回 401 Unauthorized，说明鉴权异常路径有明确响应。",
            "intro": "已补充鉴权失败场景截图，覆盖错误状态码、错误响应与后端拒绝策略证据。",
            "entries": [
                ("证据 1：Token 错误时返回 401（来源：第三阶段.pptx 第 10 页）", assets / "exception_401_response.jpg"),
                ("证据 2：Unauthorized 错误信息（来源：第三阶段.pptx 第 10 页）", assets / "exception_unauthorized_message.jpg"),
            ],
        },
    }

    for row in table.rows[1:]:
        key = row.cells[0].text.strip()
        if key not in row_data:
            continue
        set_cell_text(row.cells[0], key, size=10.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], row_data[key]["content"], size=10)
        add_evidence(row.cells[2], row_data[key]["intro"], row_data[key]["entries"])

    note = doc.add_paragraph()
    note_run = note.add_run("说明：以上截图均根据 C:\\Users\\32545\\Desktop\\软设报告 文件夹中的小程序设计草图.pdf、web端设计草图.pdf、第三阶段.pptx 整理，原始测试文档未被覆盖。")
    set_run_font(note_run, size=9)

    doc.save(str(out))
    print(out)
    print("inline_shapes", len(doc.inline_shapes))
    print("tables", len(doc.tables))


if __name__ == "__main__":
    main()
